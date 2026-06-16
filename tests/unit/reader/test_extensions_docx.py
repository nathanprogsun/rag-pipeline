"""extensions/docx 单元测试: mammoth html → html2md markdown + 内嵌图上传。

:

- Section 5.1: mammoth.convert_to_html 调用
- Section 5.2: ignore_empty_paragraphs=False 选项
- Section 5.3: 图片回调 (无 upload_file 抛错; 有 upload_file 上传 + 收集 key)
- Section 5.4: html_to_md(html, upload_file=None) — 二次 base64 不处理
- Section 5.5: mammoth 失败 → RAGError with "Can not read doc file, please convert to PDF"

覆盖:
1. minimal: 真实 sample.docx (无图) → raw_text 含文本, meta 字段完整
2. with_upload_file: mock async upload_file, 验证 docx 内嵌图被处理 (Section 5.3)
3. without_upload_file: 损坏场景下含图 + upload_file=None → 抛 RAGError (Section 5.3)
4. corrupted_buffer: 损坏 docx → RAGError with "Can not read doc file, please convert to PDF" (Section 5.5)
5. ignore_empty_paragraphs_false: mammoth 选项对齐 Section 5.2 (通过 fake 实现验证
options 传了 ignore_empty_paragraphs=False)
"""

from __future__ import annotations

import io
import struct
import zlib
from pathlib import Path
from unittest.mock import MagicMock, patch

import mammoth
import pytest

from rag.error_codes import ReaderErrorCode
from rag.exception import RAGError
from rag.ingest.reader.extensions.docx import DOCX_MIME, docx_adapter
from rag.ingest.reader.types import UploadedFileResult

# tests/data/sample.docx 路径 (与 conftest.py 的 SAMPLE_DOCX 一致)
SAMPLE_DOCX = Path(__file__).resolve().parents[2] / "data" / "sample.docx"

# 1x1 white PNG (与 test_raw_text.py / test_extensions_text.py 中的 tiny_png 同步)
_TINY_PNG_B64 = "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJAAAADUlEQVR4nGNgYGD4DwABBAEAfbLI3wAAAABJRU5ErkJggg=="


def _tiny_png_bytes() -> bytes:
    """构造 1x1 白色 PNG (避免依赖外部 fixture)。"""
    sig = b"\x89PNG\r\n\x1a\n"
    ihdr = struct.pack(">IIBBBBB", 1, 1, 8, 2, 0, 0, 0)

    def _chunk(t: bytes, d: bytes) -> bytes:
        return (
            struct.pack(">I", len(d))
            + t
            + d
            + struct.pack(">I", zlib.crc32(t + d) & 0xFFFFFFFF)
        )

    ihdr_chunk = _chunk(b"IHDR", ihdr)
    raw = b"\x00\xff\xff\xff"  # filter byte + RGB pixel
    idat_chunk = _chunk(b"IDAT", zlib.compress(raw))
    iend_chunk = _chunk(b"IEND", b"")
    return sig + ihdr_chunk + idat_chunk + iend_chunk


def _docx_with_image_bytes() -> bytes:
    """构造一份含 1 张 PNG 内嵌图的 docx (用 python-docx 写入)。"""
    from docx import Document
    from docx.shared import Inches

    png = _tiny_png_bytes()
    png_path = "/tmp/_docx_test_tiny.png"
    with open(png_path, "wb") as f:
        f.write(png)

    doc = Document()
    doc.add_paragraph("Document with embedded image")
    doc.add_picture(png_path, width=Inches(1))
    doc.add_paragraph("Tail paragraph")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── 1. minimal: 真实 sample.docx (无图) ──────────────────────────────────


@pytest.mark.asyncio
async def test_docx_extension_adapter_minimal() -> None:
    """真实 sample.docx (无图) → 走完整 mammoth → html2md 链路, raw_text 含中文段落。"""
    buf = SAMPLE_DOCX.read_bytes()
    result = await docx_adapter(buf)

    # Section 5.1 + 5.4: mammoth → html → markdown
    assert "Sample DOCX Document" in result.raw_text
    assert "Content of section A" in result.raw_text
    assert "测试 python-docx reader" in result.raw_text
    # meta 完整
    assert result.meta.mime == DOCX_MIME
    # 没图, images 必空
    assert result.images == []
    # docx 适配器只返 raw_text, 不返 format_text
    assert result.format_text is None
    assert result.extras == {}


# ── 2. with upload_file: 内嵌图上传 (Section 5.3) ───────────────────────


@pytest.mark.asyncio
async def test_docx_extension_adapter_with_upload_file() -> None:
    """内嵌图 + upload_file mock → 回调被调, key 进入 images (Section 5.3 验收)。

    用真实含图 docx (python-docx 现场生成), 跑完整链路:
    mammoth 抽到图 → _convert_image 回调 → upload_file 上传 → key 收集
    """
    buf = _docx_with_image_bytes()
    upload_calls: list[tuple[str, str, bytes]] = []

    async def upload_file(name: str, mime: str, data: bytes) -> UploadedFileResult:
        upload_calls.append((name, mime, data))
        return {"key": f"s3://bucket/{name}"}

    result = await docx_adapter(buf, upload_file=upload_file)

    # 回调被调恰好 1 次 (docx 里有 1 张图)
    assert len(upload_calls) == 1
    name, mime, data = upload_calls[0]
    # Section 5.3: name 是 uuid4 + ext_from_mime
    assert name.endswith(".png")
    assert mime == "image/png"
    assert data == _tiny_png_bytes()
    # key 进入 images
    assert result.images == [f"s3://bucket/{name}"]
    # markdown 中含图片标记 (html_to_md 输出 ![](key))
    assert f"s3://bucket/{name}" in result.raw_text
    # 段落文本仍存在
    assert "Document with embedded image" in result.raw_text
    assert "Tail paragraph" in result.raw_text


# ── 3. without upload_file + 含图 → 抛错 (Section 5.3 验收) ─────────────


@pytest.mark.asyncio
async def test_docx_extension_adapter_without_upload_file_raises() -> None:
    """含内嵌图 + upload_file=None → RAGError(READER_PARSE), message 含
    'Missing imageKeyOptions.prefix for parsed document image upload' (Section 5.3)。
    """
    buf = _docx_with_image_bytes()
    with pytest.raises(RAGError) as exc_info:
        await docx_adapter(buf, upload_file=None)
    assert exc_info.value.code == ReaderErrorCode.PARSE
    assert "Missing imageKeyOptions.prefix" in exc_info.value.message


# ── 4. corrupted buffer → RAGError (Section 5.5 验收) ────────────────────


@pytest.mark.asyncio
async def test_docx_extension_adapter_corrupted_buffer() -> None:
    """损坏 docx (非合法 zip) → RAGError(READER_PARSE), message 含
    'Can not read doc file, please convert to PDF' (Section 5.5)。"""
    with pytest.raises(RAGError) as exc_info:
        await docx_adapter(b"this is definitely not a docx")
    assert exc_info.value.code == ReaderErrorCode.PARSE
    assert "Can not read doc file, please convert to PDF" in exc_info.value.message


# ── 5. ignore_empty_paragraphs=False 选项 (Section 5.2 验收) ─────────────


@pytest.mark.asyncio
async def test_docx_extension_adapter_ignore_empty_paragraphs_false() -> None:
    """mammoth.convert_to_html 收到的 options 至少含 ``ignore_empty_paragraphs=False`` (Section 5.2)。

    用 fake mammoth 捕获 kwargs 验证选项传得对, 同时也校验无内嵌图时 upload_file=None 不抛错。
    """
    fake_result = MagicMock()
    fake_result.value = "<p>hello</p>"

    captured_kwargs: dict[str, object] = {}

    def fake_convert_to_html(*args: object, **kwargs: object) -> MagicMock:
        captured_kwargs.update(kwargs)
        return fake_result

    with patch.object(
        mammoth,
        "convert_to_html",
        side_effect=fake_convert_to_html,
    ):
        result = await docx_adapter(b"x", upload_file=None)

        # Section 5.2: ignore_empty_paragraphs 必传, 且 = False
        assert "ignore_empty_paragraphs" in captured_kwargs
        assert captured_kwargs["ignore_empty_paragraphs"] is False
        # Section 5.3: convert_image 回调总是挂载, 无 upload_file 时回调内抛错
        # (测试用 fake mammoth 不含图, 所以不会进回调, 整个链路跑通)
        assert "convert_image" in captured_kwargs
        # 跑通
        assert "hello" in result.raw_text
        assert result.images == []


# ── 附加: 错误链路保留 (cause) ─────────────────────────────────────────


@pytest.mark.asyncio
async def test_docx_extension_adapter_error_chain_preserved() -> None:
    """mammoth 抛原始异常 → RAGError 通过 ``raise ... from e`` 保留链路。"""
    raw_err = ValueError("zip corrupt inside")
    with patch.object(
        mammoth,
        "convert_to_html",
        side_effect=raw_err,
    ):
        with pytest.raises(RAGError) as exc_info:
            await docx_adapter(b"x")
        assert exc_info.value.code == ReaderErrorCode.PARSE
        # 链式原因保留
        assert exc_info.value.__cause__ is raw_err
