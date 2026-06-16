"""IngestPipeline 端到端: docx 走完整链路。"""

from __future__ import annotations

from pathlib import Path

from ingest_helpers import run_ingest
from rag.ingest.chunker import Chunker, ChunkSettings
from rag.ingest.pipeline import IngestPipeline


def test_pipeline_docx_via_file_source(sample_docx: Path) -> None:
    pipeline = IngestPipeline(
        chunker=Chunker(
            ChunkSettings(chunk_size=200, max_chunk_size=800, min_chunk_size=50)
        )
    )
    result = run_ingest(pipeline, str(sample_docx))

    assert result.chunks
    assert result.doc_meta.filename == "sample.docx"
    assert result.doc_meta.mime == (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    assert result.doc_meta.size_bytes > 0
    for c in result.chunks:
        assert c.metadata.file_type == "docx"
        assert c.metadata.encoding == "utf-8"
        assert c.metadata.source == "sample.docx"
        assert c.metadata.page_count is None


def test_pipeline_docx_chunk_metadata_has_no_heading(sample_docx: Path) -> None:
    pipeline = IngestPipeline(
        chunker=Chunker(
            ChunkSettings(chunk_size=200, max_chunk_size=800, min_chunk_size=50)
        )
    )
    result = run_ingest(pipeline, str(sample_docx))

    assert result.chunks
    for c in result.chunks:
        assert c.metadata.file_type == "docx"
        assert c.metadata.encoding == "utf-8"
        assert c.metadata.source == "sample.docx"
        assert isinstance(c.metadata.heading_stack, list)
        assert c.metadata.has_code is False
        assert c.metadata.has_table is False
    for i, c in enumerate(result.chunks):
        assert c.metadata.chunk_index == i
        assert c.metadata.total_chunks == len(result.chunks)
    full = "\n".join(c.text for c in result.chunks)
    assert "sample docx" in full.lower() or "DOCX" in full


def test_pipeline_docx_short_text_no_chunks_warning(tmp_path: Path) -> None:
    import docx as pydocx

    p = tmp_path / "tiny.docx"
    d = pydocx.Document()
    d.add_heading("Tiny", level=1)
    d.add_paragraph("One short paragraph.")
    d.save(str(p))

    pipeline = IngestPipeline(chunker=Chunker(ChunkSettings(chunk_size=200)))
    result = run_ingest(pipeline, str(p))

    assert result.chunks
    assert result.doc_meta.filename == "tiny.docx"
    assert result.doc_meta.mime == (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
